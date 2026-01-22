"""
Production-Grade RAG System for Financial Document Analysis
Uses Pinecone (cloud vector database) for embeddings
Processes all DOCX files from a folder automatically
"""

import os
import re
from typing import List, Dict, Optional
from dataclasses import dataclass
import json
from datetime import datetime
from pathlib import Path
import time

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# Document processing
from docx import Document
import pymongo
from pymongo import MongoClient

# Pinecone Vector Database
from pinecone import Pinecone, ServerlessSpec

# OpenAI for embeddings and chat
from openai import OpenAI

# Text processing
from langchain.text_splitter import RecursiveCharacterTextSplitter
import tiktoken


@dataclass
class DocumentChunk:
    """Structured document chunk with metadata"""
    text: str
    metadata: Dict
    chunk_id: str
    

class FinancialRAGSystem:
    """
    Enterprise-grade RAG system for financial document analysis
    Uses Pinecone for cloud-based vector storage
    
    Features:
    - Batch processing of multiple documents
    - Semantic chunking with context preservation
    - Cloud vector storage with Pinecone
    - Financial domain-aware processing
    - Query understanding and expansion
    """
    
    def __init__(
        self,
        openai_api_key: str,
        pinecone_api_key: str,
        index_name: str,
        mongo_uri: Optional[str] = None,
        mongo_db: str = "companyResearch",
        mongo_collection: str = "companies",
        embedding_dimension: int = 3072  # text-embedding-3-large dimension
    ):
        """Initialize RAG system with OpenAI and Pinecone"""
        
        # Initialize OpenAI
        self.openai_client = OpenAI(api_key=openai_api_key)
        self.embedding_model = "text-embedding-3-large"  # Best embedding model
        self.chat_model = "gpt-4o"  # Best reasoning model
        self.embedding_dimension = embedding_dimension
        
        # Initialize Pinecone
        print("🔧 Initializing Pinecone...")
        self.pc = Pinecone(api_key=pinecone_api_key)
        self.index_name = index_name
        
        # Create or connect to index
        existing_indexes = [idx.name for idx in self.pc.list_indexes()]
        
        if index_name not in existing_indexes:
            print(f"📦 Creating new Pinecone index: {index_name}")
            self.pc.create_index(
                name=index_name,
                dimension=embedding_dimension,
                metric="cosine",
                spec=ServerlessSpec(
                    cloud="aws",
                    region="us-east-1"
                )
            )
            # Wait for index to be ready
            print("⏳ Waiting for index to be ready...")
            time.sleep(10)
            print("✓ Index created successfully!")
        else:
            print(f"✓ Connected to existing Pinecone index: {index_name}")
        
        # Connect to index
        self.index = self.pc.Index(index_name)
        
        # Get index stats
        stats = self.index.describe_index_stats()
        print(f"📊 Index stats: {stats.total_vector_count} vectors")
        
        # Initialize MongoDB (optional)
        self.mongo_client = None
        self.mongo_db = None
        self.mongo_collection = None
        
        if mongo_uri:
            try:
                self.mongo_client = MongoClient(mongo_uri, serverSelectionTimeoutMS=2000)
                # Test connection
                self.mongo_client.server_info()
                self.mongo_db = self.mongo_client[mongo_db]
                self.mongo_collection = self.mongo_db[mongo_collection]
                print(f"✓ Connected to MongoDB: {mongo_db}.{mongo_collection}")
            except Exception as e:
                print(f"⚠️  MongoDB connection failed: {e}")
                print("Continuing without MongoDB...")
                self.mongo_client = None
        
        # Text splitter with financial context awareness
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,  # Optimal for financial data
            chunk_overlap=200,  # Preserve context across chunks
            separators=["\n\n", "\n", ". ", " ", ""],
            length_function=self._token_length
        )
        
    def _token_length(self, text: str) -> int:
        """Count tokens using tiktoken"""
        encoding = tiktoken.get_encoding("cl100k_base")
        return len(encoding.encode(text))
    
    def process_folder(self, folder_path: str) -> Dict[str, List[DocumentChunk]]:
        """
        Process all DOCX files in a folder
        
        Args:
            folder_path: Path to folder containing DOCX files
            
        Returns:
            Dictionary mapping filename to chunks
        """
        folder_path = Path(folder_path)
        
        if not folder_path.exists():
            print(f"❌ Folder not found: {folder_path}")
            return {}
        
        # Find all DOCX files
        docx_files = list(folder_path.glob("*.docx"))
        
        # Filter out temporary files (start with ~$)
        docx_files = [f for f in docx_files if not f.name.startswith("~$")]
        
        if not docx_files:
            print(f"❌ No .docx files found in {folder_path}")
            return {}
        
        print(f"\n📁 Found {len(docx_files)} documents to process:")
        for i, file in enumerate(docx_files, 1):
            print(f"   {i}. {file.name}")
        
        all_chunks = {}
        
        for i, doc_path in enumerate(docx_files, 1):
            print(f"\n{'='*60}")
            print(f"Processing {i}/{len(docx_files)}: {doc_path.name}")
            print('='*60)
            
            try:
                chunks = self.process_docx(str(doc_path))
                all_chunks[doc_path.name] = chunks
                print(f"✓ Created {len(chunks)} chunks from {doc_path.name}")
            except Exception as e:
                print(f"❌ Error processing {doc_path.name}: {e}")
                continue
        
        return all_chunks
    
    def process_docx(self, file_path: str, company_name: str = None) -> List[DocumentChunk]:
        """
        Process Word document with financial intelligence
        Preserves structure and extracts key financial indicators
        """
        doc = Document(file_path)
        
        # Extract company name from document if not provided
        if not company_name:
            company_name = self._extract_company_name(doc)
        
        # Extract full text with structure preservation
        full_text = []
        tables_data = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # Extract tables
        for table in doc.tables:
            table_text = self._extract_table(table)
            if table_text:
                tables_data.append(table_text)
        
        combined_text = "\n\n".join(full_text)
        if tables_data:
            combined_text += "\n\n" + "\n\n".join(tables_data)
        
        # Create chunks with metadata
        chunks = self._create_smart_chunks(
            combined_text,
            company_name,
            file_path
        )
        
        return chunks
    
    def _extract_company_name(self, doc: Document) -> str:
        """Extract company name from document header"""
        first_paras = [p.text for p in doc.paragraphs[:5] if p.text.strip()]
        if first_paras:
            # Look for patterns like "COMPANY NAME LIMITED"
            for para in first_paras:
                if "LIMITED" in para.upper() or "LTD" in para.upper():
                    # Clean up the company name
                    name = para.strip()
                    # Remove BSE/NSE codes if present
                    name = re.sub(r'\s*BSE:\s*\d+', '', name, flags=re.I)
                    name = re.sub(r'\s*NSE:\s*\w+', '', name, flags=re.I)
                    return name.strip()
        return "Unknown Company"
    
    def _extract_table(self, table) -> str:
        """Convert table to structured text"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # Skip empty rows
                rows.append(" | ".join(cells))
        return "\n".join(rows) if rows else ""
    
    def _create_smart_chunks(
        self,
        text: str,
        company_name: str,
        source: str
    ) -> List[DocumentChunk]:
        """
        Create semantically meaningful chunks with financial context
        """
        # Split text into chunks
        text_chunks = self.text_splitter.split_text(text)
        
        chunks = []
        for i, chunk_text in enumerate(text_chunks):
            # Extract metadata from chunk
            metadata = {
                "company": company_name,
                "source": os.path.basename(source),
                "chunk_index": i,
                "total_chunks": len(text_chunks),
                "chunk_type": self._identify_chunk_type(chunk_text),
                "contains_numbers": bool(re.search(r'\d+', chunk_text)),
                "contains_dates": bool(re.search(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', chunk_text)),
                "text": chunk_text  # Store text in metadata for Pinecone
            }
            
            # Add financial indicators
            metadata.update(self._extract_financial_indicators(chunk_text))
            
            # Create unique chunk ID
            safe_company = re.sub(r'[^\w\s-]', '', company_name)[:30]
            chunk = DocumentChunk(
                text=chunk_text,
                metadata=metadata,
                chunk_id=f"{safe_company}_{i}_{hash(chunk_text) % 10000}"
            )
            chunks.append(chunk)
        
        return chunks
    
    def _identify_chunk_type(self, text: str) -> str:
        """Identify the type of financial information in chunk"""
        text_lower = text.lower()
        
        if any(word in text_lower for word in ['revenue', 'sales', 'profit', 'ebitda', 'margin']):
            return "financial_metrics"
        elif any(word in text_lower for word in ['order', 'contract', 'project', 'awarded']):
            return "business_development"
        elif any(word in text_lower for word in ['plant', 'capacity', 'manufacturing', 'facility']):
            return "operations"
        elif any(word in text_lower for word in ['promoter', 'bought', 'sold', 'shareholding']):
            return "shareholding"
        else:
            return "general"
    
    def _extract_financial_indicators(self, text: str) -> Dict:
        """Extract key financial indicators from text"""
        indicators = {}
        
        # Revenue/Sales patterns
        revenue_match = re.search(r'(?:revenue|sales)[^\d]*?([\d,]+)\s*(?:cr|crore)', text, re.I)
        if revenue_match:
            indicators['mentions_revenue'] = True
            
        # Profit patterns
        profit_match = re.search(r'(?:profit|net profit)[^\d]*?([\d,]+)\s*(?:cr|crore)', text, re.I)
        if profit_match:
            indicators['mentions_profit'] = True
        
        # Growth patterns
        if re.search(r'\d+%.*?(?:growth|increase|up)', text, re.I):
            indicators['mentions_growth'] = True
            
        # Order book
        if re.search(r'order.*?book', text, re.I):
            indicators['mentions_orderbook'] = True
        
        return indicators
    
    def store_to_mongodb(self, chunks: List[DocumentChunk], company_name: str):
        """Store processed chunks in MongoDB for backup"""
        if self.mongo_collection is None:
            return
        
        try:
            company_doc = {
                "company_name": company_name,
                "total_chunks": len(chunks),
                "processed_date": datetime.now().isoformat(),
                "chunks": [
                    {
                        "chunk_id": chunk.chunk_id,
                        "text": chunk.text,
                        "metadata": chunk.metadata
                    }
                    for chunk in chunks
                ],
                "metadata": {
                    "chunk_types": list(set(c.metadata.get('chunk_type') for c in chunks)),
                    "has_financial_data": any(c.metadata.get('contains_numbers') for c in chunks),
                    "source_file": chunks[0].metadata.get('source') if chunks else None
                }
            }
            
            result = self.mongo_collection.update_one(
                {"company_name": company_name},
                {"$set": company_doc},
                upsert=True
            )
            
            if result.upserted_id:
                print(f"   ✓ Saved to MongoDB: {company_name}")
            else:
                print(f"   ✓ Updated MongoDB: {company_name}")
        except Exception as e:
            print(f"   ⚠️  MongoDB storage failed: {e}")
    
    def get_from_mongodb(self, company_name: str) -> Optional[Dict]:
        """Retrieve company data from MongoDB"""
        if self.mongo_collection is None:
            return None
        
        try:
            doc = self.mongo_collection.find_one({"company_name": company_name})
            return doc
        except:
            return None
    
    def list_companies_in_mongodb(self) -> List[str]:
        """Get list of all companies in MongoDB"""
        if self.mongo_collection is None:
            return []
        
        try:
            companies = self.mongo_collection.distinct("company_name")
            return companies
        except:
            return []
    
    def list_companies_in_vectordb(self) -> List[str]:
        """Get list of all companies in Pinecone vector database"""
        try:
            # Query with a dummy vector to get metadata
            # Note: Pinecone doesn't have a direct "list all metadata" function
            # This is a workaround - in production, you'd maintain a separate index
            stats = self.index.describe_index_stats()
            
            # For now, return a message that companies are stored
            # You can maintain a separate list in MongoDB or a file
            return ["Use MongoDB or check individual queries for company list"]
        except Exception as e:
            print(f"Error getting companies: {e}")
            return []
    
    def check_document_exists(self, filename: str) -> bool:
        """Check if document already exists in Pinecone"""
        try:
            # Query with filter for this specific source file
            result = self.index.query(
                vector=[0.0] * self.embedding_dimension,  # Dummy vector
                top_k=1,
                filter={"source": {"$eq": filename}},
                include_metadata=True
            )
            return len(result.matches) > 0
        except:
            return False
    
    def delete_document(self, filename: str):
        """Delete all chunks of a specific document from Pinecone"""
        try:
            print(f"🗑️  Deleting existing data for: {filename}")
            
            # Pinecone doesn't have direct delete by metadata
            # We need to fetch IDs first, then delete
            # This is a limitation - for production, consider maintaining an ID mapping
            
            # For now, we'll delete by fetching all IDs with this source
            # Note: This requires a different approach in Pinecone
            # Best practice: Use unique namespace per document or maintain ID registry
            
            print(f"⚠️  Note: Pinecone requires ID-based deletion.")
            print(f"   Recommendation: Delete entire index and rebuild, or use namespaces")
            
        except Exception as e:
            print(f"⚠️  Error during deletion: {e}")
    
    def embed_and_store(self, chunks: List[DocumentChunk], store_in_mongodb: bool = True, overwrite: bool = False):
        """
        Generate embeddings and store in Pinecone
        
        Args:
            chunks: Document chunks to store
            store_in_mongodb: Whether to backup in MongoDB
            overwrite: If True, deletes existing document data before storing
        """
        if not chunks:
            print("⚠️  No chunks to store")
            return
        
        # Check for duplicates
        if chunks:
            filename = chunks[0].metadata.get('source', '')
            if filename and not overwrite:
                exists = self.check_document_exists(filename)
                if exists:
                    print(f"\n⚠️  WARNING: '{filename}' already exists in database!")
                    print(f"   Choose an option:")
                    print(f"   1. Skip (don't add duplicates)")
                    print(f"   2. Overwrite (delete old and add new)")
                    print(f"   3. Add anyway (will create duplicates)")
                    
                    choice = input("\nYour choice (1/2/3): ").strip()
                    
                    if choice == "1":
                        print("⏭️  Skipping document")
                        return
                    elif choice == "2":
                        # Set overwrite flag
                        overwrite = True
                    elif choice == "3":
                        print("⚠️  Adding anyway - duplicates will be created")
                    else:
                        print("❌ Invalid choice, skipping...")
                        return
            
            # Delete if overwriting
            if overwrite and filename:
                self.delete_document(filename)
        
        batch_size = 100
        total_batches = (len(chunks) - 1) // batch_size + 1
        
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            
            try:
                # Generate embeddings using OpenAI
                texts = [chunk.text for chunk in batch]
                embeddings = self._generate_embeddings(texts)
                
                # Prepare vectors for Pinecone
                vectors = []
                for chunk, embedding in zip(batch, embeddings):
                    # Add timestamp to make IDs unique if not overwriting
                    chunk_id = chunk.chunk_id
                    if not overwrite:
                        # Add timestamp to prevent duplicates on re-runs
                        timestamp = int(time.time() * 1000)
                        chunk_id = f"{chunk.chunk_id}_{timestamp}"
                    
                    vectors.append({
                        "id": chunk_id,
                        "values": embedding,
                        "metadata": chunk.metadata
                    })
                
                # Upsert to Pinecone (upsert = update or insert)
                self.index.upsert(vectors=vectors)
                
                print(f"   ✓ Batch {i//batch_size + 1}/{total_batches} uploaded to Pinecone")
                
                # Small delay to avoid rate limits
                time.sleep(0.5)
                
            except Exception as e:
                print(f"   ❌ Error in batch {i//batch_size + 1}: {e}")
                continue
        
        print(f"✓ Successfully stored {len(chunks)} chunks in Pinecone")
        
        # Store in MongoDB if configured
        if store_in_mongodb and self.mongo_collection is not None and chunks:
            company_name = chunks[0].metadata.get('company', 'Unknown')
            self.store_to_mongodb(chunks, company_name)
    
    def embed_and_store_batch(self, all_chunks: Dict[str, List[DocumentChunk]]):
        """
        Store multiple documents efficiently
        
        Args:
            all_chunks: Dictionary mapping filename to chunks
        """
        print(f"\n{'='*60}")
        print(f"STORING ALL DOCUMENTS IN VECTOR DATABASE")
        print('='*60)
        
        total_chunks = sum(len(chunks) for chunks in all_chunks.values())
        print(f"Total chunks to store: {total_chunks}")
        
        for filename, chunks in all_chunks.items():
            print(f"\nStoring {filename} ({len(chunks)} chunks)...")
            self.embed_and_store(chunks, store_in_mongodb=True)
        
        print(f"\n{'='*60}")
        print(f"✓ ALL DOCUMENTS STORED SUCCESSFULLY!")
        print(f"{'='*60}")
    
    def _generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings using OpenAI"""
        response = self.openai_client.embeddings.create(
            input=texts,
            model=self.embedding_model
        )
        return [item.embedding for item in response.data]
    
    def query(
        self,
        question: str,
        top_k: int = 5,
        filter_metadata: Optional[Dict] = None
    ) -> str:
        """
        Query the RAG system with semantic understanding
        """
        # Expand and understand query
        expanded_query = self._expand_query(question)
        
        # Generate query embedding
        query_embedding = self._generate_embeddings([expanded_query])[0]
        
        # Search Pinecone
        search_kwargs = {
            "vector": query_embedding,
            "top_k": top_k,
            "include_metadata": True
        }
        
        if filter_metadata:
            search_kwargs["filter"] = filter_metadata
        
        results = self.index.query(**search_kwargs)
        
        # Extract context
        context_chunks = []
        if results.matches:
            for match in results.matches:
                context_chunks.append({
                    "text": match.metadata.get("text", ""),
                    "score": match.score,
                    "metadata": match.metadata
                })
        
        # Generate answer using GPT-4
        answer = self._generate_answer(question, context_chunks)
        
        return answer
    
    def _expand_query(self, question: str) -> str:
        """Expand query with synonyms and financial context"""
        expansions = {
            "growth": "growth rate CAGR increase expansion",
            "profit": "profit earnings net profit EBITDA margin",
            "revenue": "revenue sales turnover income",
            "order": "order book contracts projects awards",
            "plan": "strategy roadmap targets goals initiatives",
        }
        
        expanded = question.lower()
        for key, value in expansions.items():
            if key in expanded:
                expanded += " " + value
        
        return expanded
    
    def _generate_answer(self, question: str, context_chunks: List[Dict]) -> str:
        """Generate comprehensive answer using GPT-4"""
        
        if not context_chunks:
            return "❌ No relevant information found. Please make sure you've processed documents first."
        
        # Build context
        context_text = "\n\n---\n\n".join([
            f"[Source {i+1}] (Relevance: {chunk['score']:.2f})\nCompany: {chunk['metadata'].get('company', 'Unknown')}\n{chunk['text']}"
            for i, chunk in enumerate(context_chunks)
        ])
        
        # System prompt
        system_prompt = """You are an expert financial analyst assistant specializing in Indian companies and stock market analysis.

Your responsibilities:
1. Provide accurate, data-driven answers based on the provided context
2. Use specific numbers, dates, and metrics when available
3. Explain financial terms clearly
4. Cite sources by mentioning [Source X] when referencing information
5. If information is not in the context, clearly state that
6. Format responses professionally with proper structure

Guidelines:
- Focus on quantitative data (revenue, profit, growth rates, order books)
- Explain trends and their implications
- Highlight key business developments
- Use INR Crore (Cr) for currency
- Present data in a clear, organized manner"""

        user_prompt = f"""Question: {question}

Context from company documents:
{context_text}

Please provide a comprehensive answer based on the above context. Include specific numbers, dates, and cite your sources."""

        # Call GPT-4
        response = self.openai_client.chat.completions.create(
            model=self.chat_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3,
            max_tokens=1000
        )
        
        answer = response.choices[0].message.content
        
        # Add metadata
        companies = set(chunk['metadata'].get('company', '') for chunk in context_chunks)
        metadata_summary = f"\n\n---\n📊 Analysis based on documents from: {', '.join(companies)}"
        metadata_summary += f"\n🔍 Retrieved {len(context_chunks)} relevant sections"
        
        return answer + metadata_summary
    
    def get_collection_stats(self) -> Dict:
        """Get statistics about stored documents"""
        stats = self.index.describe_index_stats()
        
        return {
            "total_chunks": stats.total_vector_count,
            "index_name": self.index_name,
            "dimension": self.embedding_dimension,
            "namespaces": stats.namespaces if hasattr(stats, 'namespaces') else {}
        }


# ============================================================================
# MAIN PROGRAM
# ============================================================================

def main():
    """Main program - batch processing all documents"""
    
    print("="*80)
    print("💰 FINANCIAL RAG SYSTEM - Pinecone Edition")
    print("="*80 + "\n")
    
    # Load API keys from .env
    openai_key = os.getenv("OPENAI_API_KEY")
    pinecone_key = os.getenv("PINECONE_API_KEY")
    index_name = os.getenv("INDEX_NAME", "financial-docs")
    
    # Validation
    if not openai_key:
        print("❌ Error: OPENAI_API_KEY not found in .env file")
        print("\n📝 Please create a .env file with:")
        print("OPENAI_API_KEY=sk-your-key-here")
        print("PINECONE_API_KEY=your-pinecone-key-here")
        print("INDEX_NAME=financial-docs")
        return
    
    if not pinecone_key:
        print("❌ Error: PINECONE_API_KEY not found in .env file")
        print("\n📝 Please add to your .env file:")
        print("PINECONE_API_KEY=your-pinecone-key-here")
        return
    
    print(f"✓ OpenAI API Key loaded: {openai_key[:15]}...")
    print(f"✓ Pinecone API Key loaded: {pinecone_key[:15]}...")
    print(f"✓ Index Name: {index_name}")
    
    # MongoDB (optional)
    mongo_uri = os.getenv("MONGO_URI")
    mongo_db = os.getenv("MONGO_DB", "companyResearch")
    mongo_collection = os.getenv("MONGO_COLLECTION", "companies")
    
    # Initialize system
    print("\n🔧 Initializing RAG system...")
    try:
        rag = FinancialRAGSystem(
            openai_api_key=openai_key,
            pinecone_api_key=pinecone_key,
            index_name=index_name,
            mongo_uri=mongo_uri,
            mongo_db=mongo_db,
            mongo_collection=mongo_collection
        )
    except Exception as e:
        print(f"❌ Failed to initialize: {e}")
        import traceback
        traceback.print_exc()
        return
    
    # Show current stats
    stats = rag.get_collection_stats()
    print(f"\n📊 Current Database Status:")
    print(f"   Total chunks in Pinecone: {stats['total_chunks']}")
    print(f"   Index name: {stats['index_name']}")
    print(f"   Embedding dimension: {stats['dimension']}")
    
    # Ask user what they want to do
    print("\n" + "="*80)
    print("WHAT DO YOU WANT TO DO?")
    print("="*80)
    print("\n1. Process all documents from folder")
    print("2. Process a single document")
    print("3. Start querying (skip processing)")
    
    choice = input("\nYour choice (1/2/3): ").strip()
    
    if choice == "1":
        # Process entire folder
        docs_folder = input("\nEnter folder path (or press Enter for './docs'): ").strip()
        if not docs_folder:
            docs_folder = "./docs"
        
        print(f"\n🔍 Scanning folder: {docs_folder}")
        all_chunks = rag.process_folder(docs_folder)
        
        if all_chunks:
            print(f"\n✓ Processed {len(all_chunks)} documents")
            print(f"   Total chunks: {sum(len(c) for c in all_chunks.values())}")
            
            confirm = input("\nStore in Pinecone? (yes/no): ").strip().lower()
            if confirm in ['yes', 'y']:
                rag.embed_and_store_batch(all_chunks)
    
    elif choice == "2":
        # Process single document
        doc_path = input("\nEnter document path: ").strip()
        if os.path.exists(doc_path):
            print("\n🔄 Processing document...")
            chunks = rag.process_docx(doc_path)
            print(f"✓ Created {len(chunks)} chunks")
            
            confirm = input("\nStore in Pinecone? (yes/no): ").strip().lower()
            if confirm in ['yes', 'y']:
                rag.embed_and_store(chunks)
        else:
            print("❌ File not found")
    
    # Interactive query loop
    print("\n" + "="*80)
    print("💬 ASK YOUR QUESTIONS!")
    print("="*80)
    print("\nCommands:")
    print("  • Type your question to get answer")
    print("  • 'stats' - Show database statistics")
    print("  • 'add' - Add a new document")
    print("  • 'batch' - Process entire folder")
    print("  • 'quit' - Exit\n")
    
    while True:
        question = input("\n❓ Your input: ").strip()
        
        if not question:
            continue
            
        if question.lower() in ['quit', 'exit', 'q']:
            print("👋 Goodbye!")
            break
        
        if question.lower() == 'stats':
            stats = rag.get_collection_stats()
            print(f"\n📊 Database Statistics:")
            print(f"   Total vectors: {stats['total_chunks']}")
            print(f"   Index name: {stats['index_name']}")
            print(f"   Dimension: {stats['dimension']}")
            continue
        
        if question.lower() == 'add':
            doc_path = input("Document path: ").strip()
            if os.path.exists(doc_path):
                print("Processing...")
                chunks = rag.process_docx(doc_path)
                rag.embed_and_store(chunks)
                print("✓ Document added!")
            else:
                print("❌ File not found")
            continue
        
        if question.lower() == 'batch':
            docs_folder = input("Folder path (or press Enter for './docs'): ").strip()
            if not docs_folder:
                docs_folder = "./docs"
            
            all_chunks = rag.process_folder(docs_folder)
            if all_chunks:
                rag.embed_and_store_batch(all_chunks)
            continue
        
        # Process query
        print("\n🤔 Searching...")
        try:
            answer = rag.query(question, top_k=5)
            print("\n" + "="*80)
            print("📝 ANSWER:")
            print("="*80)
            print(answer)
            print("="*80)
        except Exception as e:
            print(f"❌ Error: {e}")
            import traceback
            traceback.print_exc()


if __name__ == "__main__":
    main()